import streamlit as st
from docx import Document
from datetime import datetime
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import uuid
import tempfile
from docx2pdf import convert
import platform

# Windows COM initialization handling
if platform.system() == "Windows":
    import pythoncom

PROPOSAL_CONFIG = {
    "Manychats + CRM Automation - 550 USD": {
        "template": "HVT Proposal - AI Automations.docx",
        "special_fields": [("VDate", "<<")],
        "team_type": "hvt_ai"
    },
    "Manychats + CRM Automation - Custom Price": {
        "template": "HVT Proposal - AI Automations - Custom Price.docx",
        "special_fields": [("VDate", "<<")],
        "team_type": "hvt_ai_custom_price"
    },
    "Internship Offer Letter": {
        "template": "Offer Letter.docx",
        "special_fields": [],
        "team_type": "offer_letter"
    }
}

def apply_run_formatting(new_run, source_run):
    """Copy formatting from source run to new run with null checks"""
    if source_run is None:
        return

    # Copy font name if it exists
    if source_run.font.name:
        new_run.font.name = source_run.font.name
        # Handle East Asian font setting safely
        rPr = new_run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), source_run.font.name)

    # Copy font size if it exists
    if source_run.font.size:
        new_run.font.size = source_run.font.size

    # Copy font color if it exists
    if source_run.font.color and source_run.font.color.rgb:
        new_run.font.color.rgb = source_run.font.color.rgb

    # Copy basic formatting properties
    new_run.bold = source_run.bold
    new_run.italic = source_run.italic
    new_run.underline = source_run.underline

def replace_placeholder(paragraph, placeholder, value):
    """Replace placeholder in paragraph while preserving formatting"""
    if placeholder not in paragraph.text:
        return False

    # Check if there are any runs
    if not paragraph.runs:
        # If no runs, just replace the text directly
        paragraph.text = paragraph.text.replace(placeholder, str(value))
        return True

    # Split runs and replace placeholder
    runs = paragraph.runs
    full_text = ''.join([run.text for run in runs])
    
    if placeholder not in full_text:
        return False

    start_idx = full_text.find(placeholder)
    end_idx = start_idx + len(placeholder)
    
    # Clear existing runs
    for run in runs:
        run.text = ""
    
    # Split text into before, replacement, and after
    before = full_text[:start_idx]
    after = full_text[end_idx:]
    
    # Add new runs with original formatting
    if before:
        new_run = paragraph.add_run(before)
        apply_run_formatting(new_run, runs[0])
    
    new_run = paragraph.add_run(str(value))
    apply_run_formatting(new_run, runs[0])
    
    if after:
        new_run = paragraph.add_run(after)
        apply_run_formatting(new_run, runs[-1])
    
    return True

def process_document(doc, placeholders):
    """Process entire document including tables and nested tables"""
    # Process paragraphs
    for paragraph in doc.paragraphs:
        if not paragraph.text:
            continue  # Skip empty paragraphs
        for ph, value in placeholders.items():
            replace_placeholder(paragraph, ph, value)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Handle nested tables first
                if cell.tables:
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                for para in nested_cell.paragraphs:
                                    if not para.text:
                                        continue
                                    for ph, value in placeholders.items():
                                        replace_placeholder(para, ph, value)
                # Process cell paragraphs
                for para in cell.paragraphs:
                    if not para.text:
                        continue
                    for ph, value in placeholders.items():
                        replace_placeholder(para, ph, value)
                # Maintain cell alignment
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return doc

def get_hvt_ai_team_details():
    """Collect team composition details with improved layout"""
    st.subheader("Team Composition")
    team_roles = {
        "Project Manager": "P1",
        "Frontend Developers": "F1",
        "UI/UX Members": "U1",
        "AI/ML Developers": "A1",
        "Business Analyst": "B1",
        "AWS Developer": "AD1",
        "Backend Developers": "BD1",
        "System Architect": "S1"
    }
    
    team_details = {}
    cols = st.columns(2)
    
    for idx, (role, placeholder) in enumerate(team_roles.items()):
        with cols[idx % 2]:
            team_details[f"<<{placeholder}>>"] = st.number_input(
                f"{role} Count:",
                min_value=0,
                step=1,
                key=f"hvt_team_{placeholder}"
            )
    return team_details

def validate_phone_number(country, number):
    """Enhanced phone number validation"""
    if not number:
        return True
    if country.lower() == "india":
        return number.startswith("+91")
    return number.startswith("+1")

def generate_document():
    st.title("Document Generator Pro")
    base_dir = os.path.join(os.getcwd(), "templates")

    selected_proposal = st.selectbox("Select Document Type", list(PROPOSAL_CONFIG.keys()))
    config = PROPOSAL_CONFIG[selected_proposal]
    template_path = os.path.join(base_dir, config["template"])

    # Initialize session state for downloads
    if 'generated_files' not in st.session_state:
        st.session_state.generated_files = {}

    # Collect common fields
    placeholders = {}
    if selected_proposal == "Internship Offer Letter":
        st.subheader("Candidate Information")
        placeholders.update({
            "<<E-Name>>": st.text_input("Candidate Name:"),
            "<<Job>>": st.selectbox("Job Role", ["UI UX", "AI Automations", "Software Developer", "Sales"]),
            "<<S-Date>>": st.date_input("Start Date").strftime("%d %B, %Y"),
            "<<Stipend>>": f"{st.number_input('Stipend (₹)', min_value=0):,}",
            "<<Months>>": st.number_input("Duration (Months)", min_value=1),
            "<<Date>>": datetime.today().strftime("%d %B, %Y")
        })
    else:
        st.subheader("Client Details")
        col1, col2 = st.columns(2)
        with col1:
            client_name = st.text_input("Client Name:")
            client_email = st.text_input("Email:")
        with col2:
            country = st.text_input("Country:")
            client_number = st.text_input("Phone Number:")
        
        # Add separate date fields
        st.subheader("Date Information")
        date_col1, date_col2 = st.columns(2)
        with date_col1:
            date_field = st.date_input("Proposal Date", datetime.today())
        with date_col2:
            validation_date = st.date_input("Validation Date", datetime.today())

        placeholders.update({
            "<<Client Name>>": client_name,
            "<<Client Email>>": client_email,
            "<<Client Number>>": client_number,
            "<<Country>>": country,
            "<<Date>>": date_field.strftime("%d %B, %Y"),
            "<<D-Date>>": date_field.strftime("%d %B, %Y"),  # Same as Date
            "<<VDate>>": validation_date.strftime("%d-%m-%Y")  # New validation date
        })

        # Add team composition
        if "hvt_ai" in config["team_type"]:
            placeholders.update(get_hvt_ai_team_details())

        # Add pricing section
        if "custom_price" in config["team_type"]:
            st.subheader("Pricing Details")
            pricing = {
                "<<P01>>": st.number_input("Manychats Setup (USD)", min_value=0),
                "<<P02>>": st.number_input("Make Automations (USD)", min_value=0),
                "<<A-Price>>": st.number_input("Annual Maintenance (USD)", min_value=0)
            }
            placeholders.update(pricing)
            placeholders["<<T-Price>>"] = f"{sum(pricing.values()):,}"

    if st.button("Generate Documents"):
        # Validate inputs
        if selected_proposal != "Internship Offer Letter":
            if not validate_phone_number(placeholders["<<Country>>"], placeholders["<<Client Number>>"]):
                st.error("Invalid phone number format for selected country")
                return

        # Generate unique filenames
        unique_id = uuid.uuid4().hex[:8]
        base_name = f"{selected_proposal.replace(' ', '_')}_{unique_id}"
        doc_filename = f"{base_name}.docx"
        pdf_filename = f"{base_name}.pdf"

        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                # Process Word document
                doc = Document(template_path)
                doc = process_document(doc, placeholders)
                doc_path = os.path.join(temp_dir, doc_filename)
                doc.save(doc_path)

                # Convert to PDF
                pdf_path = os.path.join(temp_dir, pdf_filename)
                if platform.system() == "Windows":
                    pythoncom.CoInitialize()
                convert(doc_path, pdf_path)
                if platform.system() == "Windows":
                    pythoncom.CoUninitialize()

                # Store files in session state
                with open(doc_path, "rb") as f:
                    st.session_state.generated_files['doc'] = f.read()
                with open(pdf_path, "rb") as f:
                    st.session_state.generated_files['pdf'] = f.read()
                st.session_state.generated_files['doc_name'] = doc_filename
                st.session_state.generated_files['pdf_name'] = pdf_filename

                st.success("Documents generated successfully!")

        except Exception as e:
            st.error(f"Generation failed: {str(e)}")
            if platform.system() == "Windows":
                pythoncom.CoUninitialize()

    # Display download buttons
    if 'doc' in st.session_state.generated_files:
        st.markdown("---")
        st.subheader("Download Generated Files")
        
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="Download Word Document",
                data=st.session_state.generated_files['doc'],
                file_name=st.session_state.generated_files['doc_name'],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with col2:
            st.download_button(
                label="Download PDF Document",
                data=st.session_state.generated_files['pdf'],
                file_name=st.session_state.generated_files['pdf_name'],
                mime="application/pdf"
            )

if __name__ == "__main__":
    generate_document()
